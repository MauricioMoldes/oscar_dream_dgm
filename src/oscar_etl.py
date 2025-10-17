
"""
===========================================================
OSCAR-DREAM Genomic Report Parser
===========================================================

Description:
-------------
This Python script parses sensitive genomic reports in DOCX format.
It extracts sample information, test details, gene variants,
and clinical interpretations, storing the data in a structured format
(e.g., pandas DataFrame or CSV).

Intended for internal use only within the OSCAR-DREAM project unit.
All sample data must be handled in compliance with GDPR and internal
data privacy protocols.

Maintainer:
-------------
Name: Mauricio Moldes
Email: mauricio.moldes.quaresma@regionh.dk
Affiliation: Unit for Genomic Medicine

Version:
--------
1.0.0

Creation Date: 2025-10-08
Last Modified: 2025-10-08

Python Version:
---------------
>= 3.10

Dependencies:
-------------
- python-docx
- pandas
- regex
- nltk / spacy (optional for NLP)
- hashlib (for anonymization)

Usage:
------
$ python genomic_report_parser.py --input "report.docx" --output "parsed.csv"

Notes:
------
- Ensure that all DOCX files are stored and accessed securely.
- Sample identifiers should be anonymized if data leaves secure environment.
- Extend regex patterns and NLP models for additional structured data extraction.
===========================================================
"""

import argparse
import pandas as pd
from docx import Document
import hashlib
import re
import sys
import os


# ----------------------------
# Helper Functions
# ----------------------------


def extract_report_version(filename):
    """
    Extracts the version number (e.g., V3, V4, V5) from a DOCX filename.

    Parameters:
        filename (str): Path or name of the DOCX file.

    Returns:
        str: Version string (e.g., "V3"), or None if not found.
    """
    # Extract only the filename part in case a full path is passed
    base_name = os.path.basename(filename)

    # Search for version patterns like V3, v3, _V5, -V10 etc.
    match = re.search(r'[ _-]V(\d+)', base_name, re.IGNORECASE)
    if match:
        return f"V{match.group(1)}"
    else:
        return None
    


def dump_docx(file_path):
    """
    Dump all readable text content from a DOCX file, including:
    - Paragraphs
    - Tables
    - Header/footer text (if available)

    Parameters:
        file_path (str): Path to DOCX file.

    Returns:
        str: Combined plain text from the document.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error opening file {file_path}: {e}")
        return None

    full_text = []

    # --- Extract paragraphs ---
    full_text.append("=== PARAGRAPHS ===")
    for para in doc.paragraphs:
        full_text.append(para.text)

    # --- Extract tables ---
    full_text.append("\n=== TABLES ===")
    for table_idx, table in enumerate(doc.tables, start=1):
        full_text.append(f"\n--- Table {table_idx} ---")
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            full_text.append(f"Row {row_idx}: " + " | ".join(cells))

    # --- Extract headers & footers (if any) ---
    full_text.append("\n=== HEADERS & FOOTERS ===")
    for section_idx, section in enumerate(doc.sections, start=1):
        header = section.header
        footer = section.footer
        if header and header.paragraphs:
            full_text.append(f"\n--- Section {section_idx} Header ---")
            for para in header.paragraphs:
                full_text.append(para.text)
        if footer and footer.paragraphs:
            full_text.append(f"\n--- Section {section_idx} Footer ---")
            for para in footer.paragraphs:
                full_text.append(para.text)

    return "\n".join(full_text)


def read_docx(file_path):
    """
    Reads a DOCX file and returns the text as a single string.
    
    Parameters:
        file_path (str): Path to the DOCX file.
    
    Returns:
        str: Full text of the document.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error opening file {file_path}: {e}")
        return None

    # Extract paragraphs that are not empty
    full_text = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    
    return "\n".join(full_text)


def anonymize_id(sample_id):
    """Anonymize sample identifiers using SHA-256 hashing."""
    return hashlib.sha256(sample_id.encode()).hexdigest()

def parse_sample_info(text):
    """Placeholder: Extract sample information from text."""
    sample_info = {}
    # Example: extract DOB
    dob_match = re.search(r"Date of Birth[:\s]+(\d{2}/\d{2}/\d{4})", text)
    sample_info['DOB'] = dob_match.group(1) if dob_match else None
    # Add other fields like sample ID, sample ID
    return sample_info

def parse_variants(text):
    """Placeholder: Extract gene variants and classifications."""
    variants = []
    # Example structure: [{"Gene": "BRCA1", "Variant": "c.68_69delAG", "Classification": "Pathogenic"}]
    return variants

# ----------------------------
# Version-specific Parsers
# ----------------------------

def parse_v1(file_path):
    print("Parsing using V1 logic")
    # TODO: implement version-specific parsing
    return dump_docx(file_path)

def parse_v2(file_path):
    print("Parsing using V2 logic")
    return dump_docx(file_path)

def parse_v3(file_path):
    print("Parsing using V3 logic")
    return dump_docx(file_path)

def parse_v4(file_path):
    print("Parsing using V4 logic")
    return dump_docx(file_path)

def parse_v5(file_path):
    print("Parsing using V5 logic")
    return dump_docx(file_path)

def parse_v6(file_path):
    print("Parsing using V6 logic")
    return dump_docx(file_path)

def parse_vn_plus_1(file_path):
    print("Using fallback parser for unknown or new version")
    return dump_docx(file_path)


# ----------------------------
# Version Dispatcher (Auxiliary Function)
# ----------------------------

def dispatch_parser_by_version(file_path):
    """
    Auxiliary function that selects which parser to use
    based on the detected report version.
    """
    version = extract_report_version(file_path)
    print(f"=== DOCUMENT VERSION === {version}")

    match version:
        case "V1":
            return parse_v1(file_path)
        case "V2":
            return parse_v2(file_path)
        case "V3":
            return parse_v3(file_path)
        case "V4":
            return parse_v4(file_path)
        case "V5":
            return parse_v5(file_path)
        case "V6":
            return parse_v6(file_path)
        case _:
            return parse_vn_plus_1(file_path)



# ----------------------------
# Main Function
# ----------------------------

def main():
    """parser = argparse.ArgumentParser(description="Parse genomic DOCX reports into structured CSV")
    parser.add_argument("--input", required=True, help="Path to input DOCX report")
    parser.add_argument("--output", required=True, help="Path to output CSV file")
    args = parser.parse_args()"""

    # working examples for V1-> V6 
   
    file_path="dev" 
     
    content = dispatch_parser_by_version(file_path)

    if content:
        print("\n=== RAW DOCUMENT DUMP ===")
        print(content)
    else:
        print("Failed to read the document.")

    #text = read_docx(file_path)
    
"""
    if text:
        print("=== Document Content Preview ===")
        print(text) 
        #print(text[:1000]) # preview first 1000 characters
    else:
        print("Failed to read the document.")
"""
"""
    # Step 1: Read DOCX
    try:
        text = read_docx(args.input)
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        sys.exit(1)

    # Step 2: Parse sample info
     sample_info = parse_sample_info(text)
    
    # Step 3: Anonymize sample ID (if exists)
    if "sample_ID" in sample_info and sample_info["sample_ID"]:
        sample_info["sample_ID"] = anonymize_id(sample_info["sample_ID"])

    # Step 4: Parse gene variants
    variants = parse_variants(text)

    # Step 5: Combine data into a DataFrame
    df = pd.DataFrame(variants)
    for key, value in sample_info.items():
        df[key] = value

    # Step 6: Save to CSV
    try:
        df.to_csv(args.output, index=False)
        print(f"Parsed report saved to {args.output}")
    except Exception as e:
        print(f"Error saving CSV file: {e}")
        sys.exit(1)"""

# ----------------------------
# Entry Point
# ----------------------------

if __name__ == "__main__":
    main()

